import datetime
import shutil
import smtplib
import unittest
from datetime import time
from email.mime.image import MIMEImage
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.utils import formataddr

import xlrd
from docx import Document
from docx.shared import Inches
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from BeautifulReport import BeautifulReport
from common.all_path import *


# =================初始化信息=================
test_dir = case_path  # 测试用例代码路径 #指定执行脚本的目录
test_report = report_path  # 指定存放报告的目录
# print(test_dir,test_report)


# 企业邮箱账号
# my_sender = 'guoqi.dong@geneseeq.com'
my_sender = 'guanzhong.zhou@geneseeq.com'
# 发件人邮箱密码(是我企业邮箱授权密码)
# my_pass = 'SsgNh26eQ2BXqQFB'
my_pass = '7gEdadjuD2bPZzdd'

send_date = (datetime.datetime.now()).strftime('%Y-%m-%d')
# print(send_date)
# test_report = report_path #指定存放报告的目录


# ========全量执行用例【前置方法】==========
def find_pyfile_and_import():
    """遍历用例下的test。py文件，全部import成功"""
    casedir = case_path  # 指定了用例路径
    # 遍历文件夹找出的指定类型文件，导入，注意globals,否则作用域只是在这个函数下
    for dirName, subdirList, fileList in os.walk(casedir):
        print(dirName)
        for f in fileList:
            file_name = f
            if file_name[0:5] == 'test_' and file_name[-3:] == '.py':
                impath = (dirName.split("Lims_Test")[1])
                # print(impath)
                impath = impath[1:].replace("\\", ".")
                # print(impath)

                if impath != "":
                    # exe_str = "from " + impath + " import " + file_name[0:-3]
                    exe_str = "from " + impath + "." + file_name[0:-3] + " import " + "*"
                    print(exe_str)
                    exec(exe_str, globals())


def get_xls_case_by_index(sheet_name):
    """
    获取excel表格里指定的sheet数据，保存到列表返回
    :param sheet_name: sheet表名
    :return:
    """

    global col1, col3, col2
    xls_path = case_xls_path  # 指定了用例表格文件路径
    file = xlrd.open_workbook(xls_path)
    sheet = file.sheet_by_name(sheet_name)
    ncols = sheet.ncols
    for j in range(ncols):
        cell_value = sheet.cell_value(0, j)
        if cell_value == "fileName":
            col1 = j
        elif cell_value == "ClassName":
            col2 = j
        elif cell_value == "caseName":
            col3 = j
    nrows = sheet.nrows
    caseList = []

    for i in range(1, nrows):
        # 注意这里的'y'不是'Y'
        if sheet.row_values(i)[0].lower().strip() == 'y':
            # print(sheet.row_values(i))
            ClassName = sheet.cell_value(i, col2)
            # 组装测试用例名称格式：文件名.类名('方法名')
            # case='%s.%s("%s")' % (fileName.strip(),ClassName.strip(),caseName.strip())
            case = '%s' % (ClassName.strip())
            caseList.append(case)
            # print(case)
    # print(caseList)
    return caseList


# =========全量执行用例===========
def Run_allcase():
    """
    执行全量用例
    """
    # 运行前先清空文件夹下过期的图片
    shutil.rmtree(img_path)
    os.mkdir(img_path)

    # 创建一个测试套件
    suite = unittest.TestSuite()
    # 创建一个用例加载对象
    loader = unittest.TestLoader()

    find_pyfile_and_import()  # py文件导入成功
    testCaseList = get_xls_case_by_index('Sheet1')  # 将要加载到测试套的用例，预加载在list中(这里是将类加载到测试套，方法加载到测试套一直
    print(testCaseList)
    # 报错，目前定位不出问题)

    # 注意这里执行excel表格的用例，最低维度是执行class类下的所有case，所以打标签最低是类上
    for test_case in testCaseList:
        suite.addTest(loader.loadTestsFromTestCase(eval(test_case)))

    nowTime = time.strftime("%Y-%m-%d_%H-%M-%S")
    testName = 'Lims系统自动化'
    testResult = BeautifulReport(suite)
    testResult.report(filename=str(testName) + str('_') + str(nowTime), description='用例执行情况', log_path=test_report)

    # 列出目录的下所有文件和文件夹保存到lists
    lists = os.listdir(test_report)
    # 按时间排序
    lists.sort(key=lambda fn: os.path.getmtime(test_report + "\\" + fn))
    # 获取最新的文件保存到file_new
    file_new = os.path.join(test_report, lists[-1])


# =========用例执行完自动发送邮件【前置方法】===========
def send_mail(*parse):
    """使用QQ邮箱/企业微信邮箱发送邮件"""
    ret = True
    # noinspection PyBroadException
    try:
        # ************************准备附件&图片********************
        # 读取附件
        f1 = open(parse[0], 'rb')
        mail_body = f1.read()
        f1.close()

        # 读取图片
        f2 = open(parse[4], "rb")
        img_data = f2.read()
        f2.close()

        # ************************定义图片属性********************
        msg = MIMEMultipart('related')
        content = MIMEText('<html><body><img src="cid:imageid" alt="imageid"></body></html>', 'html', 'utf-8')
        msg.attach(content)
        img = MIMEImage(img_data)
        img.add_header('Content-ID', 'imageid')
        msg.attach(img)  # 将图片加入

        # ************************定义附件属性********************
        # 附件1，传送上一个函数得出的最新测试报告文件
        att = MIMEText(mail_body, 'html', 'utf-8')
        att["Content-Type"] = 'application/octet-stream'
        # 这里的filename可以任意写,Test_Reports.html,这里要指定.html格式，不然接受者打不开
        att["Content-Disposition"] = 'attachment; filename=' + \
                                     str(send_date) + '_Test_Reports.html'
        # 附件2
        # ....
        # 附件N
        msg.attach(att)  # 将附件加入

        # ************************定义收件人，发送人，抄送人信息********************
        to_addrs = parse[1].split(',')  # 主送人+多人
        acc_to_addrs = parse[2].split(',')  # 抄送人+多人

        # 括号里的对应发件人邮箱昵称，发件人邮箱账号
        msg['From'] = formataddr(["发送报告机器人-世小和", my_sender])

        # 括号里的对应收件人邮箱昵称，收件人邮箱账号
        msg['To'] = ",".join(to_addrs)

        # 抄送
        msg['Cc'] = ",".join(acc_to_addrs)

        # 邮件的主题，也可以说是标题
        msg['Subject'] = parse[3]

        # 发件人邮箱中的SMTP服务器，端口是25
        server = smtplib.SMTP_SSL("smtp.exmail.qq.com", 465)  # 使用企业邮箱发送
        # 括号中对应的是发件人邮箱账号，邮箱密码
        server.login(my_sender, my_pass)
        # 括号中对应的是发件人邮箱账号，收件人邮箱账号，发送邮件
        server.sendmail(my_sender, to_addrs + acc_to_addrs,
                        msg.as_string())  # 如果有抄送人，要把主送人和抄送人+链接

        # 关闭连接
        server.quit()
    # 如果 try 中的语句没有执行，则会执行下面的 ret=False
    except Exception:
        ret = False
    return ret


def new_report():
    # ***********************找报告的html文件**************************
    # 列出目录的下所有文件和文件夹保存到lists
    lists = os.listdir(test_report)
    # 按时间排序
    lists.sort(key=lambda fn: os.path.getmtime(test_report + "\\" + fn))
    # 获取最新的文件保存到file_new
    file_new = os.path.join(test_report, lists[-1])

    return file_new


def get_image(pic_name):
    # ***********************对报告文件截屏**************************
    # 截屏开始时间
    start = time.time()
    # 得到html文件路径
    file_new = new_report()
    # 设置chrome开启的模式，headless就是无界面模式
    # 一定要使用这个模式，不然截不了全页面，只能截到你电脑的高度
    chrome_options = Options()
    chrome_options.add_argument('headless')
    driver = webdriver.Chrome(chrome_options=chrome_options)

    # 控制浏览器写入并转到链接
    # 这里的url因为我的文件是本地的html所以需要你需要动态的网页你自己编写一下
    driver.get(file_new)

    time.sleep(1)
    # 接下来是全屏的关键，用js获取页面的宽高，如果有其他需要用js的部分也可以用这个方法
    width = driver.execute_script("return document.documentElement.scrollWidth")
    height = driver.execute_script("return document.documentElement.scrollHeight")
    print("当前输出报告的html长宽为：" + str(width) + "&" + str(height))
    # 将浏览器的宽高设置成刚刚获取的宽高
    driver.set_window_size(width, height)
    # driver.set_window_size(1000, height)
    driver.set_window_size(1929, 4230)
    time.sleep(1)
    # 截图并关掉浏览器
    driver.save_screenshot(pic_name)
    driver.close()
    print('HTML---->png转换完成，花费时间：', time.time() - start)

    return file_new, pic_name


# =========自动发送邮件===========
def Auto_sendmail():
    """
    自动发送邮件
    """
    # os.system('python MyRun.py')#执行MyRun文件
    # 收件人邮箱账号,逗号隔开，发送多人邮件
    # 收件人（'guoqi.dong@geneseeq.com,...,xxxx@163.com'）
    receive_user = 'guoqi.dong@geneseeq.com'
    # receive_user = 'guoqi.dong@geneseeq.com,guanzhong.zhou@geneseeq.com,benjamin.yan@geneseeq.com'
    # 抄送人（'1923180362@qq.com,...,xxxx@163.com'）
    acc = 'guanzhong.zhou@geneseeq.com'
    # acc = 'yongxiang.wang@geneseeq.com,xiaojun.cheng@geneseeq.com,yanshou.gao@geneseeq.com,qingjun.zhang@geneseeq.com,' \
    # 'wei.hu@geneseeq.com,yuanbin.tang@geneseeq.com'
    email_title = "您有一份世和LIMS系统的自动化测试报告，请查收"  # 邮件标题
    # 调用取最新报告文件的截图方法
    get_report = get_image(img_path + '\\' + 'report_shot.png')
    report_file = get_report[0]  # 获取html文件
    report_png = get_report[1]  # 获取png文件

    # 调用发送邮件方法
    ret = send_mail(report_file, receive_user, acc, email_title, report_png)
    # 判断邮件是否发送成功，给执行者打印信息
    if ret:
        print("Send Success!")
    else:
        print("Send Failure!")


# =========自动截图文件汇总========
def Creat_auto_shotfile():
    """
    自动将报告运行中的截图形成文件并说明
    """

    # 遍历文件夹找出的指定类型文件，导入，注意globals,否则作用域只是在这个函数下
    # list_dirs = os.walk(img_path)
    fileList, filenamedesc, fileimg = [], [], []
    doc = Document()
    num = 0
    for dirName, subdirList, filename in os.walk(img_path):

        fileList.append(filename)
        for f in filename:
            num += 1
            desc = f.split('_')[-1].split('.')[0]  # 分离出步骤说明文字
            filenamedesc.append(str(num) + '.' + desc)
            pic_path = os.path.join(os.path.join(img_path, f))  # 获得每张图片的路径
            fileimg.append(pic_path)
            print(pic_path)

    with open(fileshot, "a") as file:
        for i in range(len(filenamedesc)):
            print(filenamedesc[i])
            # file.write(filenamedesc[i]+'\n'+'\n')  # 添加步骤描述
            doc.add_paragraph().add_run(str(filenamedesc[i])).bold = True  # 正常添加文字，但是到word中文字会有部分加粗了
            doc.add_picture(fileimg[i], width=Inches(6), height=Inches(3))  # 添加图片  # 添加图片
            doc.save(fileshot)


if __name__ == '__main__':
    # 运行前先清空文件夹下过期的图片
    shutil.rmtree(img_path)
    os.mkdir(img_path)
    now = time.strftime("%Y-%m-%d_%H-%M-%S")
    name = str(now) + '_report'
    test_suite = unittest.defaultTestLoader.discover(
        test_dir, pattern='test0*.py')
    result = BeautifulReport(test_suite)
    result.report(filename=name, description='Lims系统自动化-用例执行情况',
                  log_path=test_report)
    Creat_auto_shotfile()
